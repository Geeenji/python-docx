# -*- coding: utf-8 -*-
#
# test_parts.py
#
# Copyright (C) 2013 Steve Canny scanny@cisco.com
#
# This module is part of python-docx and is released under the MIT License:
# http://www.opensource.org/licenses/mit-license.php

"""Test suite for the docx.parts module."""

from docx.parts import _Body, _Document

import pytest

from mock import call, Mock

from .unitdata import a_body
from .unitutil import class_mock, function_mock, initializer_mock


class Describe_Document(object):

    @pytest.fixture
    def _Body_(self, request):
        return class_mock('docx.parts._Body', request)

    @pytest.fixture
    def init(self, request):
        return initializer_mock(_Document, request)

    @pytest.fixture
    def oxml_fromstring_(self, request):
        return function_mock('docx.parts.oxml_fromstring', request)

    @pytest.fixture
    def oxml_tostring(self, request):
        return function_mock('docx.parts.oxml_tostring', request)

    def it_can_be_constructed_by_opc_part_factory(
            self, oxml_fromstring_, init):
        # mockery ----------------------
        partname, content_type, blob, document_elm = (
            Mock(name='partname'), Mock(name='content_type'),
            Mock(name='blob'), Mock(name='document_elm')
        )
        oxml_fromstring_.return_value = document_elm
        # exercise ---------------------
        doc = _Document.load(partname, content_type, blob)
        # verify -----------------------
        oxml_fromstring_.assert_called_once_with(blob)
        init.assert_called_once_with(partname, content_type, document_elm)
        assert isinstance(doc, _Document)

    def it_has_a_body(self, init, _Body_):
        # mockery ----------------------
        doc = _Document(None, None, None)
        doc._element = Mock(name='_element')
        # exercise ---------------------
        body = doc.body
        # verify -----------------------
        _Body_.assert_called_once_with(doc._element.body)
        assert body is _Body_.return_value

    def it_can_serialize_to_xml(self, init, oxml_tostring):
        # mockery ----------------------
        doc = _Document(None, None, None)
        doc._element = Mock(name='_element')
        # exercise ---------------------
        doc.blob
        # verify -----------------------
        oxml_tostring.assert_called_once_with(
            doc._element, encoding='UTF-8', pretty_print=False,
            standalone=True)


class Describe_Body(object):

    @pytest.fixture
    def Paragraph_(self, request):
        return class_mock('docx.parts.Paragraph', request)

    def it_can_add_a_paragraph_to_itself(self, Paragraph_):
        # mockery ----------------------
        body_elm = Mock(name='body_elm')
        body_elm.add_p.return_value = p_elm = Mock(name='p_elm')
        body = _Body(body_elm)
        # exercise ---------------------
        p = body.add_paragraph()
        # verify -----------------------
        body_elm.add_p.assert_called_once_with()
        Paragraph_.assert_called_once_with(p_elm)
        assert p is Paragraph_.return_value

    def it_provides_access_to_its_paragraphs_as_a_sequence(self, Paragraph_):
        # mockery ----------------------
        body_elm = Mock(name='body_elm')
        p1, p2 = (Mock(name='p1'), Mock(name='p2'))
        P1, P2 = (Mock(name='Paragraph1'), Mock(name='Paragraph2'))
        body_elm.p = [p1, p2]
        body = _Body(body_elm)
        Paragraph_.side_effect = [P1, P2]
        # exercise ---------------------
        paragraphs = body.paragraphs
        # verify -----------------------
        assert Paragraph_.mock_calls == [call(p1), call(p2)]
        assert paragraphs == (P1, P2)

    def it_returns_an_empty_sequence_when_it_contains_no_paragraphs(self):
        body = _Body(a_body().element)
        assert body.paragraphs == ()

    def it_can_clear_itself_of_all_content_it_holds(self):
        # mockery ----------------------
        body_elm = Mock(name='body_elm')
        body = _Body(body_elm)
        # exercise ---------------------
        retval = body.clear_content()
        # verify -----------------------
        body_elm.clear_content.assert_called_once_with()
        assert retval is body
